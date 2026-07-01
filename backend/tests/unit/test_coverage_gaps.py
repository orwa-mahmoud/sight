"""Tests targeting remaining coverage gaps across the codebase.

Covers:
- WhatsApp webhook: verify invalid UUID, full POST flow with mocked deps, _send_reply, _check_signature
- Agent loop: max iterations reached, unknown tool dispatch
- Checkpoint: conversation not found, empty messages after threshold, LLM exception, empty summary text
- Chunker: hard token split edge cases (empty separator, _token_chunks)
- Parser: DOCX parsing
- Graph: execute_tools non-AI message, save_key_fact dispatch, escalate_question dispatch
"""

from __future__ import annotations

import hashlib
import hmac as hmac_mod
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import pytest

from src.ai.context.checkpoint import maybe_create_checkpoint
from src.domain.conversations.value_objects import ConversationChannel
from src.domain.llm.value_objects import LLMCallResult, TokenUsage
from src.infrastructure.channels.whatsapp import WhatsAppAdapter

# ── WhatsApp webhook ──────────────────────────────────────────────


class TestWhatsAppVerifyInvalidUUID:
    """Verify endpoint with invalid UUID returns 400."""

    async def test_verify_bad_uuid_returns_400(self) -> None:
        from httpx import ASGITransport, AsyncClient

        from src.main import app

        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            resp = await client.get(
                "/webhooks/not-a-uuid/whatsapp",
                params={"hub.mode": "subscribe", "hub.verify_token": "token", "hub.challenge": "ch"},
            )
        assert resp.status_code == 400


class TestWhatsAppAdapterVerifySignature:
    """Verify WhatsAppAdapter.verify_signature with per-tenant app_secret."""

    def test_verify_signature_valid(self) -> None:
        body = b'{"test": "data"}'
        sig = "sha256=" + hmac_mod.new(b"app-secret", body, hashlib.sha256).hexdigest()
        assert WhatsAppAdapter.verify_signature(body, sig, "app-secret") is True

    def test_verify_signature_invalid(self) -> None:
        assert WhatsAppAdapter.verify_signature(b"body", "sha256=bad", "app-secret") is False


class TestWhatsAppWebhookPostFlow:
    """Cover lines 67-95: full POST handler with a text message, mocked deps."""

    async def test_webhook_post_config_not_found(self) -> None:
        """When tenant config is None, return 404."""
        from httpx import ASGITransport, AsyncClient

        from src.main import app

        tid = uuid4()
        payload = {
            "entry": [
                {
                    "changes": [
                        {
                            "value": {
                                "messages": [{"from": "+971501234567", "type": "text", "text": {"body": "hello"}}],
                                "contacts": [{"profile": {"name": "Test"}}],
                                "metadata": {"phone_number_id": "pn123"},
                            }
                        }
                    ]
                }
            ]
        }
        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            resp = await client.post(f"/webhooks/{tid}/whatsapp", json=payload)
        # Config doesn't exist for random UUID -> 404
        assert resp.status_code == 404

    @patch("src.infrastructure.channels.whatsapp.WhatsAppAdapter.send_text", new_callable=AsyncMock)
    @patch("src.drivers.api.webhooks.whatsapp.chat_with_agent", new_callable=AsyncMock)
    async def test_webhook_post_success_with_reply(self, mock_chat: AsyncMock, mock_send: AsyncMock) -> None:
        """Full success flow: config found, signature passes, agent called, reply sent."""
        from src.application.shared.unit_of_work import UnitOfWork
        from src.domain.tenant_config.entities import TenantConfig
        from src.domain.tenants.entities import Tenant
        from src.domain.users.entities import User, UserTenant
        from src.domain.users.value_objects import UserTenantRole
        from src.infrastructure.persistence.postgres.database import async_session_factory

        # Setup: create tenant + config with WA tokens + app_secret
        async with async_session_factory() as session:
            uow = UnitOfWork(session)
            tenant = Tenant.create(name="WA Test", slug=f"wa-{uuid4().hex[:8]}")
            await uow.tenants.save(tenant)
            user = User.create(email=f"{uuid4().hex[:8]}@test.com", hashed_password="hash")
            await uow.users.save(user)
            await uow.flush()
            link = UserTenant.create(user_id=user.id, tenant_id=tenant.id, role=UserTenantRole.OWNER)
            await uow.user_tenants.save(link)
            config = TenantConfig.create_default(tenant_id=tenant.id)
            config._is_new = True
            config.update_whatsapp(
                phone_number_id="pn123",
                access_token="EAA-secret",
                verify_token="",
                app_secret="test-app-secret",
            )
            await uow.tenant_configs.save(config)
            await uow.commit()
            tid = tenant.id

        # Mock chat_with_agent to return a ChatResult
        from src.ai.types import ChatResult

        mock_chat.return_value = ChatResult(
            response="Hello from bot!",
            thread_id="test-thread",
        )

        import json

        from httpx import ASGITransport, AsyncClient

        from src.main import app

        payload = {
            "entry": [
                {
                    "changes": [
                        {
                            "value": {
                                "messages": [{"from": "+971501234567", "type": "text", "text": {"body": "hello"}}],
                                "contacts": [{"profile": {"name": "Test"}}],
                                "metadata": {"phone_number_id": "pn123"},
                            }
                        }
                    ]
                }
            ]
        }
        body = json.dumps(payload).encode()
        sig = "sha256=" + hmac_mod.new(b"test-app-secret", body, hashlib.sha256).hexdigest()

        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            resp = await client.post(
                f"/webhooks/{tid}/whatsapp",
                content=body,
                headers={"Content-Type": "application/json", "X-Hub-Signature-256": sig},
            )

        assert resp.status_code == 200
        mock_chat.assert_called_once()
        mock_send.assert_called_once()

    @patch("src.drivers.api.webhooks.whatsapp.chat_with_agent", new_callable=AsyncMock)
    async def test_webhook_post_agent_exception_returns_503(self, mock_chat: AsyncMock) -> None:
        """When chat_with_agent raises, the handler asks Meta to redeliver (503)
        instead of acking a lost reply."""
        from src.application.shared.unit_of_work import UnitOfWork
        from src.domain.tenant_config.entities import TenantConfig
        from src.domain.tenants.entities import Tenant
        from src.domain.users.entities import User, UserTenant
        from src.domain.users.value_objects import UserTenantRole
        from src.infrastructure.persistence.postgres.database import async_session_factory

        async with async_session_factory() as session:
            uow = UnitOfWork(session)
            tenant = Tenant.create(name="WA Err", slug=f"wae-{uuid4().hex[:8]}")
            await uow.tenants.save(tenant)
            user = User.create(email=f"{uuid4().hex[:8]}@test.com", hashed_password="hash")
            await uow.users.save(user)
            await uow.flush()
            link = UserTenant.create(user_id=user.id, tenant_id=tenant.id, role=UserTenantRole.OWNER)
            await uow.user_tenants.save(link)
            config = TenantConfig.create_default(tenant_id=tenant.id)
            config._is_new = True
            config.update_whatsapp(app_secret="err-secret")
            await uow.tenant_configs.save(config)
            await uow.commit()
            tid = tenant.id

        mock_chat.side_effect = RuntimeError("boom")

        import json

        from httpx import ASGITransport, AsyncClient

        from src.main import app

        payload = {
            "entry": [
                {
                    "changes": [
                        {
                            "value": {
                                "messages": [{"from": "+123", "type": "text", "text": {"body": "hi"}}],
                                "contacts": [{"profile": {"name": "X"}}],
                                "metadata": {"phone_number_id": "pn1"},
                            }
                        }
                    ]
                }
            ]
        }
        body = json.dumps(payload).encode()
        sig = "sha256=" + hmac_mod.new(b"err-secret", body, hashlib.sha256).hexdigest()

        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            resp = await client.post(
                f"/webhooks/{tid}/whatsapp",
                content=body,
                headers={"Content-Type": "application/json", "X-Hub-Signature-256": sig},
            )

        assert resp.status_code == 503


class TestWhatsAppVerifySuccess:
    """Cover line 46: successful verification returns the challenge."""

    async def test_verify_success_returns_challenge(self) -> None:
        from src.application.shared.unit_of_work import UnitOfWork
        from src.domain.tenant_config.entities import TenantConfig
        from src.domain.tenants.entities import Tenant
        from src.domain.users.entities import User, UserTenant
        from src.domain.users.value_objects import UserTenantRole
        from src.infrastructure.persistence.postgres.database import async_session_factory

        async with async_session_factory() as session:
            uow = UnitOfWork(session)
            tenant = Tenant.create(name="VerTest", slug=f"vt-{uuid4().hex[:8]}")
            await uow.tenants.save(tenant)
            user = User.create(email=f"{uuid4().hex[:8]}@test.com", hashed_password="hash")
            await uow.users.save(user)
            await uow.flush()
            link = UserTenant.create(user_id=user.id, tenant_id=tenant.id, role=UserTenantRole.OWNER)
            await uow.user_tenants.save(link)
            config = TenantConfig.create_default(tenant_id=tenant.id)
            config._is_new = True
            config.update_whatsapp(verify_token="my-verify-token")
            await uow.tenant_configs.save(config)
            await uow.commit()
            tid = tenant.id

        from httpx import ASGITransport, AsyncClient

        from src.main import app

        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            resp = await client.get(
                f"/webhooks/{tid}/whatsapp",
                params={
                    "hub.mode": "subscribe",
                    "hub.verify_token": "my-verify-token",
                    "hub.challenge": "challenge-string-123",
                },
            )

        assert resp.status_code == 200
        assert resp.text == "challenge-string-123"


class TestWhatsAppAdapterSendText:
    """WhatsAppAdapter.send_text — success and error paths."""

    async def test_send_text_no_credentials(self) -> None:
        adapter = WhatsAppAdapter(phone_number_id="", access_token="")
        result = await adapter.send_text("+123", "Hello!")
        assert result is None


class TestWhatsAppWebhookSignatureCheck:
    """Cover line 67-75: POST handler with valid signature check (has verify token)."""

    async def test_webhook_post_bad_signature_returns_403(self) -> None:
        """When signature check fails, return 403."""
        from src.application.shared.unit_of_work import UnitOfWork
        from src.domain.tenant_config.entities import TenantConfig
        from src.domain.tenants.entities import Tenant
        from src.domain.users.entities import User, UserTenant
        from src.domain.users.value_objects import UserTenantRole
        from src.infrastructure.persistence.postgres.database import async_session_factory

        async with async_session_factory() as session:
            uow = UnitOfWork(session)
            tenant = Tenant.create(name="SigTest", slug=f"sig-{uuid4().hex[:8]}")
            await uow.tenants.save(tenant)
            user = User.create(email=f"{uuid4().hex[:8]}@test.com", hashed_password="hash")
            await uow.users.save(user)
            await uow.flush()
            link = UserTenant.create(user_id=user.id, tenant_id=tenant.id, role=UserTenantRole.OWNER)
            await uow.user_tenants.save(link)
            config = TenantConfig.create_default(tenant_id=tenant.id)
            config._is_new = True
            config.update_whatsapp(
                access_token="EAA-token",
                app_secret="sig-test-secret",
            )
            await uow.tenant_configs.save(config)
            await uow.commit()
            tid = tenant.id

        from httpx import ASGITransport, AsyncClient

        from src.main import app

        payload = {
            "entry": [
                {
                    "changes": [
                        {
                            "value": {
                                "messages": [{"from": "+123", "type": "text", "text": {"body": "hi"}}],
                                "contacts": [],
                                "metadata": {"phone_number_id": "pn1"},
                            }
                        }
                    ]
                }
            ]
        }

        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            resp = await client.post(
                f"/webhooks/{tid}/whatsapp",
                json=payload,
                headers={"X-Hub-Signature-256": "sha256=wrong"},
            )

        assert resp.status_code == 403


# ── Agent loop ────────────────────────────────────────────────────


# ── Checkpoint ────────────────────────────────────────────────────


class TestCheckpointEdgeCases:
    """Cover lines 77, 87, 99-101, 105 in checkpoint.py."""

    async def test_conversation_not_found_returns_none(self) -> None:
        """Line 77: conv not found -> early return."""
        mock_uow = MagicMock()
        mock_uow.conversations = MagicMock()
        mock_uow.conversations.get_by_thread_id = AsyncMock(return_value=None)
        mock_llm = AsyncMock()

        await maybe_create_checkpoint(
            thread_id="nonexistent",
            tenant_id=uuid4(),
            channel=ConversationChannel.WEB,
            llm=mock_llm,
            uow=mock_uow,
        )
        mock_llm.chat_with_tools.assert_not_called()

    async def test_empty_messages_returns_none(self) -> None:
        """Line 87: messages list is empty after threshold check -> early return."""
        mock_conv = MagicMock()
        mock_conv.id = uuid4()

        mock_uow = MagicMock()
        mock_uow.conversations = MagicMock()
        mock_uow.conversations.get_by_thread_id = AsyncMock(return_value=mock_conv)
        mock_uow.messages = MagicMock()
        mock_uow.messages.sum_tokens_since_checkpoint = AsyncMock(return_value=5000)
        mock_uow.messages.list_since_last_checkpoint = AsyncMock(return_value=[])
        mock_llm = AsyncMock()

        await maybe_create_checkpoint(
            thread_id="test-thread",
            tenant_id=uuid4(),
            channel=ConversationChannel.WEB,
            llm=mock_llm,
            uow=mock_uow,
        )
        mock_llm.chat_with_tools.assert_not_called()

    async def test_llm_exception_returns_none(self) -> None:
        """Lines 99-101: LLM raises exception -> caught, returns None."""
        mock_conv = MagicMock()
        mock_conv.id = uuid4()

        mock_msg = MagicMock()
        mock_msg.is_checkpoint = False
        mock_msg.role = "user"
        mock_msg.content = "test message"

        mock_uow = MagicMock()
        mock_uow.conversations = MagicMock()
        mock_uow.conversations.get_by_thread_id = AsyncMock(return_value=mock_conv)
        mock_uow.messages = MagicMock()
        mock_uow.messages.sum_tokens_since_checkpoint = AsyncMock(return_value=5000)
        mock_uow.messages.list_since_last_checkpoint = AsyncMock(return_value=[mock_msg])

        mock_llm = AsyncMock()
        mock_llm.chat_with_tools.side_effect = RuntimeError("LLM down")

        await maybe_create_checkpoint(
            thread_id="test-thread",
            tenant_id=uuid4(),
            channel=ConversationChannel.WEB,
            llm=mock_llm,
            uow=mock_uow,
        )
        # Should not raise, just log warning

    async def test_empty_summary_returns_none(self) -> None:
        """Line 105: LLM returns empty summary text -> early return."""
        mock_conv = MagicMock()
        mock_conv.id = uuid4()

        mock_msg = MagicMock()
        mock_msg.is_checkpoint = False
        mock_msg.role = "user"
        mock_msg.content = "test"

        mock_uow = MagicMock()
        mock_uow.conversations = MagicMock()
        mock_uow.conversations.get_by_thread_id = AsyncMock(return_value=mock_conv)
        mock_uow.messages = MagicMock()
        mock_uow.messages.sum_tokens_since_checkpoint = AsyncMock(return_value=5000)
        mock_uow.messages.list_since_last_checkpoint = AsyncMock(return_value=[mock_msg])
        mock_uow.token_usages = MagicMock()
        mock_uow.token_usages.save = AsyncMock()
        mock_uow.track = MagicMock()

        mock_llm = AsyncMock()
        mock_llm.chat_with_tools.return_value = LLMCallResult(
            text="   ",  # whitespace only -> empty after strip
            usage=TokenUsage(input_tokens=10, output_tokens=5),
            provider="openai",
            model="gpt-4o-mini",
        )

        await maybe_create_checkpoint(
            thread_id="test-thread",
            tenant_id=uuid4(),
            channel=ConversationChannel.WEB,
            llm=mock_llm,
            uow=mock_uow,
        )
        # No checkpoint message is saved, but the billable summarizer call is still
        # recorded (an empty reply still cost input+output tokens).
        mock_uow.token_usages.save.assert_awaited_once()


# ── Chunker edge cases ────────────────────────────────────────────


class TestChunkerEdgeCases:
    """Cover lines 50, 53, 65-66: hard token split when text is huge without separators."""

    def test_token_chunks_direct(self) -> None:
        """Lines 65-66: _token_chunks splits text into exact token-sized pieces."""
        from src.infrastructure.rag.chunker import RecursiveTokenChunker

        chunker = RecursiveTokenChunker(chunk_size=10)
        # A string with no separators that is long enough to require _token_chunks
        result = chunker._token_chunks("abcdefghijklmnopqrstuvwxyz " * 10)
        assert len(result) > 1

    def test_no_separators_triggers_token_split(self) -> None:
        """Lines 50, 53: when separators exhausted, falls through to _token_chunks."""
        from src.infrastructure.rag.chunker import RecursiveTokenChunker

        # Use a tiny chunk size so a long continuous string without separators triggers the path
        chunker = RecursiveTokenChunker(chunk_size=5)
        text = "a" * 200  # No whitespace, no separators at all
        chunks = chunker.chunk(text)
        assert len(chunks) > 1


# ── Parser edge cases ─────────────────────────────────────────────


class TestParserDocx:
    """Cover lines 42-43: DOCX parsing (both valid and invalid)."""

    def test_parse_docx_invalid_content_raises(self) -> None:
        from src.domain.documents.value_objects import DocumentMimeType
        from src.domain.shared.exceptions import InvalidOperationError
        from src.infrastructure.rag.parser import parse

        with pytest.raises(InvalidOperationError, match="Could not parse DOCX"):
            parse(b"not a valid docx file", DocumentMimeType.DOCX)

    def test_parse_docx_valid(self) -> None:
        """Lines 42-43: successful DOCX parsing returns joined paragraphs."""
        import io

        from docx import Document as DocxDocument

        from src.domain.documents.value_objects import DocumentMimeType
        from src.infrastructure.rag.parser import parse

        doc = DocxDocument()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        result = parse(buf.read(), DocumentMimeType.DOCX)
        assert "First paragraph" in result
        assert "Second paragraph" in result


# ── Graph execution edges ─────────────────────────────────────────


class TestGraphDispatchEdges:
    """Cover lines 84, 201, 211 in graph.py."""

    async def test_execute_tools_non_ai_message(self) -> None:
        """Line 84: execute_tools with non-AI last message returns state unchanged.

        We build the graph and then invoke the execute_tools node directly
        with a state whose last message is HumanMessage (not AIMessage).
        """
        from langchain_core.messages import HumanMessage

        from src.infrastructure.ai.graph import build_agent_graph

        mock_llm = AsyncMock()
        mock_llm.chat_with_tools = AsyncMock(
            return_value=LLMCallResult(
                text="Direct answer",
                usage=TokenUsage(input_tokens=10, output_tokens=5),
            )
        )
        graph = build_agent_graph(
            llm=mock_llm,
            tools=[],
            retriever=AsyncMock(),
            uow=AsyncMock(),
        )
        # Get the async node function from the graph — the node spec wraps a RunnableCallable
        node_spec = graph.nodes["execute_tools"]
        execute_tools_fn = node_spec.runnable.afunc  # type: ignore[union-attr]
        state = {
            "messages": [HumanMessage(content="hi")],
            "tool_calls_made": [],
            "total_input_tokens": 0,
            "total_output_tokens": 0,
            "total_cache_tokens": 0,
            "iteration": 0,
            "tenant_id": str(uuid4()),
            "channel": "web",
            "conversation_id": None,
            "contact_id": None,
        }
        result = await execute_tools_fn(state)
        # Should return state unchanged since last message is not AIMessage
        assert result["messages"] == state["messages"]

    async def test_dispatch_escalate_question(self) -> None:
        """Line 201: escalate_question dispatch in graph."""
        from src.infrastructure.ai.graph import _dispatch_tool

        with patch("src.infrastructure.ai.graph.run_escalate_question", new_callable=AsyncMock) as mock_esc:
            mock_esc.return_value = {"status": "submitted"}
            result = await _dispatch_tool(
                tool_name="escalate_question",
                arguments={"question": "test"},
                tenant_id=uuid4(),
                channel=ConversationChannel.WEB,
                conversation_id=uuid4(),
                contact_id=uuid4(),
                retriever=AsyncMock(),
                uow=MagicMock(),
            )
            assert result["status"] == "submitted"

    async def test_dispatch_save_key_fact(self) -> None:
        """Line 211: save_key_fact dispatch in graph."""
        from src.infrastructure.ai.graph import _dispatch_tool

        with patch("src.infrastructure.ai.graph.run_save_key_fact", new_callable=AsyncMock) as mock_skf:
            mock_skf.return_value = {"status": "saved"}
            result = await _dispatch_tool(
                tool_name="save_key_fact",
                arguments={"key": "name", "value": "Alice"},
                tenant_id=uuid4(),
                channel=ConversationChannel.WEB,
                conversation_id=None,
                contact_id=uuid4(),
                retriever=AsyncMock(),
                uow=MagicMock(),
            )
            assert result["status"] == "saved"
